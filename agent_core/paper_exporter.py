"""Markdown 论文 → DOCX 转换器，支持内嵌图片。

功能:
    - 解析 Markdown 中的 ![alt](path) 图片引用
    - 将图片从本地路径嵌入到 DOCX 文档中
    - 处理标题、粗体、斜体等基础 Markdown 格式
    - 输出可独立分发的带内嵌图片的 Word 文档
"""

import os
import re
import zipfile
from pathlib import Path
from typing import Optional
from io import BytesIO

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

_BR_L = chr(0x5B)
_BR_R = chr(0x5D)
_PA_L = chr(0x28)
_PA_R = chr(0x29)

_MATH_PATTERN = re.compile(r"\x24\x24([\s\S]*?)\x24\x24")

def _render_latex_as_image(latex: str, dpi: int = 150) -> Optional[BytesIO]:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    try:
        fig, ax = plt.subplots(figsize=(6, 0.5))
        ax.axis("off")
        ax.text(0.5, 0.5, f"${latex}$", fontsize=13, ha="center", va="center",
                transform=ax.transAxes)
        buf = BytesIO()
        fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight",
                    facecolor="white", edgecolor="none")
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception:
        return None

IMG_PATTERN = re.compile(
    "!\\" + _BR_L + "([^\\" + _BR_R + "]*)\\" + _BR_R
    + "\\" + _PA_L + "([^" + _PA_R + "]+)\\" + _PA_R
)


def _resolve_image_path(img_ref: str, base_dir: Path) -> Optional[Path]:
    cleaned = img_ref.strip()
    # Candidate 1: exact path relative to base_dir
    candidate = base_dir / cleaned
    if candidate.is_file() and candidate.suffix.lower() in (".png", ".jpg", ".jpeg", ".gif", ".bmp"):
        return candidate
    # Candidate 2: if model wrote "workspace/charts/xxx.png", resolve under base_dir
    if cleaned.startswith("workspace/"):
        candidate = base_dir / cleaned[len("workspace/"):]
        if candidate.is_file() and candidate.suffix.lower() in (".png", ".jpg", ".jpeg", ".gif", ".bmp"):
            return candidate
    # Candidate 3: search by filename in charts/ and workspace/charts/
    fname = Path(cleaned).name
    for search_dir in [
        base_dir / "workspace" / "charts",
        base_dir / "charts",
        base_dir / "workspace" / "data",
    ]:
        candidate = search_dir / fname
        if candidate.is_file() and candidate.suffix.lower() in (".png", ".jpg", ".jpeg", ".gif", ".bmp"):
            return candidate
    # Candidate 4: glob search for the filename
    for search_root in [base_dir / "workspace", base_dir]:
        if search_root.exists():
            for found in search_root.rglob(fname):
                if found.is_file():
                    return found
    return None


def _add_formatted_text_to_paragraph(p, text: str):
    p.style.font.size = Pt(11)
    pattern = re.compile(
        r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)"
    )
    last_end = 0
    for m in pattern.finditer(text):
        if m.start() > last_end:
            run = p.add_run(text[last_end:m.start()])
            run.font.size = Pt(11)
        if m.group(2):
            run = p.add_run(m.group(2))
            run.bold = True
            run.font.size = Pt(11)
        elif m.group(3):
            run = p.add_run(m.group(3))
            run.italic = True
            run.font.size = Pt(11)
        elif m.group(4):
            run = p.add_run(m.group(4))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        last_end = m.end()
    if last_end < len(text):
        run = p.add_run(text[last_end:])
        run.font.size = Pt(11)


def _set_heading_font(heading, size: int):
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "\u9ed1\u4f53")


def _insert_image(doc, img_path: Path, alt_text: str, max_width: float):
    from PIL import Image as PILImage
    with PILImage.open(img_path) as pil_img:
        orig_w, orig_h = pil_img.size
    dpi = 150
    inches_w = orig_w / dpi
    inches_h = orig_h / dpi
    usable_width = max_width - 1.0
    if inches_w > usable_width:
        scale = usable_width / inches_w
        inches_w = usable_width
        inches_h *= scale
    last_paragraph = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = last_paragraph.add_run()
    run.add_picture(str(img_path), width=Inches(inches_w), height=Inches(inches_h))
    if alt_text:
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption.add_run(alt_text)
        cap_run.font.size = Pt(9)
        cap_run.italic = True
        cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _process_markdown_block(doc, block: str):
    lines = block.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("---"):
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(qn("w:bottom"), {
                qn("w:val"): "single", qn("w:sz"): "6",
                qn("w:space"): "1", qn("w:color"): "999999",
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            continue
        if stripped.startswith("#### "):
            h = doc.add_heading(stripped[5:], level=4)
            _set_heading_font(h, 12)
        elif stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=3)
            _set_heading_font(h, 13)
        elif stripped.startswith("## "):
            h = doc.add_heading(stripped[2:], level=2)
            _set_heading_font(h, 14)
        elif stripped.startswith("# "):
            h = doc.add_heading(stripped[2:], level=1)
            _set_heading_font(h, 16)
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(qn("w:bottom"), {
                qn("w:val"): "single", qn("w:sz"): "12",
                qn("w:space"): "4", qn("w:color"): "333333",
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text_to_paragraph(p, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text_to_paragraph(p, re.sub(r"^\d+\.\s", "", stripped))
        else:
            _add_formatted_text_to_paragraph(doc.add_paragraph(), stripped)


def _fix_docx_xml(docx_path: Path):
    """Strip XML declarations from internal XML files in a docx archive.

    python-docx may write <?xml...?> at the start of word/document.xml and other
    internal XML files. Word (especially older versions) rejects such files.
    This re-packs the archive with clean XML.
    """
    import tempfile
    import shutil

    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(tmp_fd)
    try:
        with zipfile.ZipFile(str(docx_path), 'r') as zin:
            with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename.endswith(('.xml', '.rels')):
                        text = data.decode('utf-8', errors='replace')
                        # Strip leading XML declaration (<?xml ...?>)
                        text = re.sub(r'^<\?xml\s+[^?]*\?\s*>', '', text, count=1)
                        # Strip any leading whitespace left behind (Word rejects it)
                        text = text.lstrip()
                        data = text.encode('utf-8')
                    zout.writestr(item, data)
        shutil.move(tmp_path, str(docx_path))
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def md_to_docx(
    md_content: str,
    output_path: Path,
    base_dir: Optional[Path] = None,
    max_image_width: float = 14.0,
) -> dict:
    if not HAS_DOCX:
        return {"status": "skipped", "error": "python-docx not installed"}
    if base_dir is None:
        base_dir = Path.cwd()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    images_embedded = 0
    images_missed = 0

    math_dir = output_path.parent / ".math_cache"
    math_dir.mkdir(exist_ok=True)
    math_idx = [0]

    def _replace_math(match):
        latex = match.group(1).strip()
        buf = _render_latex_as_image(latex)
        if buf:
            math_idx[0] += 1
            img_path = math_dir / f"math_{math_idx[0]:04d}.png"
            img_path.write_bytes(buf.read())
            return "\n![公式](" + str(img_path).replace("\\", "/") + ")\n"
        return "\n[公式渲染失败]\n"

    md_content = _MATH_PATTERN.sub(_replace_math, md_content)

    segments = IMG_PATTERN.split(md_content)
    i = 0
    while i < len(segments):
        if i + 2 < len(segments):
            text_segment = segments[i]
            alt_text = segments[i + 1]
            img_ref = segments[i + 2]
            if text_segment:
                _process_markdown_block(doc, text_segment)
            img_path = _resolve_image_path(img_ref, base_dir)
            if img_path:
                try:
                    _insert_image(doc, img_path, alt_text, max_image_width)
                    images_embedded += 1
                except Exception:
                    p = doc.add_paragraph()
                    run = p.add_run(f"[image: {alt_text or img_ref}]")
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x99, 0x33, 0x33)
                    images_missed += 1
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"[image not found: {alt_text or img_ref}]")
                run.italic = True
                run.font.color.rgb = RGBColor(0x99, 0x33, 0x33)
                images_missed += 1
            i += 3
        else:
            _process_markdown_block(doc, segments[i])
            i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    # Fix: python-docx may write XML declaration in document.xml which breaks Word.
    # Re-pack the docx, stripping XML declarations from all internal XML files.
    _fix_docx_xml(output_path)

    return {
        "status": "generated",
        "output_path": str(output_path),
        "images_embedded": images_embedded,
        "images_missed": images_missed,
    }


def export_existing_paper(md_path: Path, output_path: Optional[Path] = None) -> dict:
    if not md_path.exists():
        return {"status": "error", "error": f"file not found: {md_path}"}
    md_content = md_path.read_text(encoding="utf-8")
    base_dir = md_path.resolve().parent.parent.parent
    if output_path is None:
        output_path = md_path.with_suffix(".docx")
    return md_to_docx(md_content, output_path, base_dir=base_dir)
